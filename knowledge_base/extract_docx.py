#!/usr/bin/env python3
"""Konwertuje wszystkie DOCX do plików Markdown w knowledge_base/docs/"""
import os, re, json
from pathlib import Path
from docx import Document

SRC = Path("/home/user/webapp/knowledge_base/raw")
DST = Path("/home/user/webapp/knowledge_base/docs")
DST.mkdir(parents=True, exist_ok=True)

def slugify(name: str) -> str:
    s = name.lower()
    s = re.sub(r'\.docx$', '', s)
    s = re.sub(r'[^a-z0-9]+', '-', s)
    s = re.sub(r'-+', '-', s).strip('-')
    return s[:80]

def docx_to_md(path: Path) -> str:
    doc = Document(path)
    out = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            out.append("")
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if "heading 1" in style:
            out.append(f"# {text}")
        elif "heading 2" in style:
            out.append(f"## {text}")
        elif "heading 3" in style:
            out.append(f"### {text}")
        elif "heading 4" in style:
            out.append(f"#### {text}")
        elif "list" in style or text.startswith(("•", "·", "-", "*")):
            out.append(f"- {re.sub(r'^[•·\\-\\*]\\s*', '', text)}")
        else:
            out.append(text)
    # Tables
    for table in doc.tables:
        out.append("")
        for i, row in enumerate(table.rows):
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            out.append("| " + " | ".join(cells) + " |")
            if i == 0:
                out.append("| " + " | ".join(["---"] * len(cells)) + " |")
        out.append("")
    return "\n".join(out)

manifest = []
for docx in sorted(SRC.glob("*.docx")):
    try:
        md = docx_to_md(docx)
    except Exception as e:
        print(f"ERR {docx.name}: {e}")
        continue
    slug = slugify(docx.name)
    out_path = DST / f"{slug}.md"
    out_path.write_text(md, encoding="utf-8")
    manifest.append({
        "source": docx.name,
        "slug": slug,
        "md_path": str(out_path),
        "chars": len(md),
        "lines": md.count("\n") + 1
    })
    print(f"OK  {docx.name} -> {out_path.name}  ({len(md)} chars)")

(DST.parent / "manifest.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")
print(f"\nTotal: {len(manifest)} files, {sum(m['chars'] for m in manifest):,} chars")
